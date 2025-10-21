# apps/imports/views.py
from zipfile import ZipFile, BadZipFile
from io import BytesIO
import re, json, os
from datetime import date, datetime
from django.utils.dateparse import parse_date
from django.conf import settings
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.response import Response
from rest_framework import status
from django.db import transaction
from django.shortcuts import get_object_or_404
from typing import Union, List, Optional

from apps.users.models import OfficerProfile
from apps.directory.models import Rank, Unit
from core.permissions import IsAdminOrRoot, IsHR
from docx import Document

EMAIL_RE = re.compile(r'[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}')

MONTHS = {
    # и рус, и англ на всякий случай
    'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
    'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12,
    'января': 1, 'февраля': 2, 'марта': 3, 'апреля': 4, 'мая': 5, 'июня': 6,
    'июля': 7, 'августа': 8, 'сентября': 9, 'октября': 10, 'ноября': 11, 'декабря': 12,
}


def _norm_month_date(text: str):
    """
    '25 September 1995' -> '1995-09-25'
    '03 November 1994' -> '1994-11-03'
    '09.2014' -> '2014-09-01'
    '01.08.2023' -> '2023-08-01'
    'по н/время' -> None
    """
    t = (text or "").strip()
    if not t or 'н/время' in t.lower():
        return None

    # dd.mm.yyyy -> yyyy-mm-dd
    m = re.match(r'^(\d{2})\.(\d{2})\.(\d{4})$', t)
    if m:
        d, mon, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        return f"{y:04d}-{mon:02d}-{d:02d}"

    # dd Month yyyy
    m = re.match(r'(\d{1,2})\s+([A-Za-zА-Яа-яё]+)\s+(\d{4})', t)
    if m:
        d, mon, y = int(m.group(1)), m.group(2).lower(), int(m.group(3))
        mon = MONTHS.get(mon, None)
        if mon:
            return f"{y:04d}-{mon:02d}-{d:02d}"
    # mm.yyyy
    m = re.match(r'(\d{2})\.(\d{4})', t)
    if m:
        mon, y = int(m.group(1)), int(m.group(2))
        return f"{y:04d}-{mon:02d}-01"
    # yyyy-mm-dd
    d = parse_date(t)
    return d.isoformat() if d else None


def _find_email(lines: list[str]) -> Optional[str]:
    """
    Ищем email в строках документа.
    Сначала по префиксам: 'E-mail', 'email', 'почта', 'эл. почта', затем — любое совпадение regexp.
    """
    prefixes = ("e-mail", "email", "почта", "эл. почта", "эл.почта", "электронная почта")
    for ln in lines:
        low = ln.lower()
        if any(low.startswith(p + ":") or low.startswith(p + " :") for p in prefixes):
            # после двоеточия парсим email
            after = ln.split(":", 1)[-1]
            m = EMAIL_RE.search(after)
            if m:
                return m.group(0).strip()
    # если явного поля нет — ищем просто любой email в тексте
    for ln in lines:
        m = EMAIL_RE.search(ln)
        if m:
            return m.group(0).strip()
    return None

def _map_marital_status(text: str) -> str:
    """
    Приводим строку семейного положения к choices модели:
    SINGLE / MARRIED / DIVORCED / WIDOWED
    """
    t = (text or "").strip().lower()

    if not t:
        return ""

    # позитивные/частые варианты
    if any(x in t for x in ["женат", "замужем", "брак", "супруг", "супруга"]):
        return "MARRIED"
    if any(x in t for x in ["холост", "не замужем", "одинок"]):
        return "SINGLE"
    if any(x in t for x in ["развед", "расторг", "бывш"]):
        return "DIVORCED"
    if any(x in t for x in ["вдов", "вдова", "вдовец"]):
        return "WIDOWED"

    # иначе не трогаем (пускай вручную поправят)
    return ""


def parse_ld8_docx(file_bytes: bytes, filename: str) -> dict:
    doc = Document(BytesIO(file_bytes))

    # 1) Собираем линейный текст (параграфы)
    para_lines = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            para_lines.append(t)

    # Быстрый геттер по префиксу в параграфах
    def find_after(prefix):
        for ln in para_lines:
            if ln.lower().startswith(prefix.lower()):
                return ln.split(":", 1)[-1].strip()
        return ""

    # 2) Звание + дата в скобках: "Капитан (01.08.2023)"
    rank_line_idx = -1
    rank_name, rank_since = None, None
    for i, ln in enumerate(para_lines):
        m = re.match(r'^([А-ЯЁA-Z][а-яёa-zA-ZЁ\- ]+)\s*\((\d{2}\.\d{2}\.\d{4})\)$', ln)
        if m:
            rank_line_idx = i
            rank_name = m.group(1).strip()
            rank_since = _norm_month_date(m.group(2))
            break

    # 3) ФИО — следующая непустая строка после строки со званием
    full_name = ""
    if rank_line_idx != -1:
        j = rank_line_idx + 1
        while j < len(para_lines) and (not para_lines[j].strip() or para_lines[j].strip().upper() == "СПРАВКА"):
            j += 1
        if j < len(para_lines):
            if re.match(r'^[А-ЯЁA-Z][а-яёa-zЁ\-]+ [А-ЯЁA-Z][а-яёa-zЁ\-]+(?: [А-ЯЁA-Z][а-яёa-zЁ\-]+)?$', para_lines[j]):
                full_name = para_lines[j].strip()

    if not full_name:
        cand = find_after("Ф.И.О.")
        if cand:
            full_name = cand

    # 4) Остальные реквизиты из параграфов
    personal_number = find_after("личный номер")
    birth_line = find_after("Число, месяц, год и место рождения")
    # '25 September 1995 года, г. Атырау'
    birth_date, birth_place = None, None
    if birth_line:
        # разделим по 'года,' или последней запятой
        parts = [p.strip() for p in re.split(r'года,|,', birth_line, maxsplit=1) if p.strip()]
        if parts:
            birth_date = _norm_month_date(re.sub(r'\s*года$', '', parts[0]))
        if len(parts) > 1:
            birth_place = parts[1].strip()

    email = _find_email(para_lines)

    payload = {
        "source_file": filename,
        "rank": {"name": rank_name, "since": rank_since} if rank_name else None,
        "full_name": full_name,
        "personal_number": personal_number,
        "birth": {"date": birth_date, "place": birth_place} if birth_date or birth_place else None,
        "iin": find_after("Индивидуальный идентификационный номер"),
        "nationality": find_after("Национальность"),
        "education": {
            "civil": find_after("а) гражданское"),
            "military": find_after("б) военное"),
        },
        "awards": find_after("Государственные награды"),
        "penalties": find_after("Взыскания"),
        "combat_participation": find_after("Участие в боевых действиях"),
        "foreign_trips": find_after("Длительные заграничные командировки"),
        "marital_status": find_after("Семейное положение"),
        "email": email,
        "service_history": [],
        "sign_block": {}
    }

    # 5) История службы: сначала пытаемся из ТАБЛИЦЫ
    def _clean(s: str) -> str:
        return " ".join((s or "").split())

    def _norm_cell_date(s: str):
        s = (s or "").strip()
        if not s:
            return None
        # "по н/время" -> None
        if "н/время" in s.lower():
            return None
        # mm.yyyy -> yyyy-mm-01
        if re.match(r'^\d{2}\.\d{4}$', s):
            return _norm_month_date(s)
        # dd.mm.yyyy -> yyyy-mm-dd
        if re.match(r'^\d{2}\.\d{2}\.\d{4}$', s):
            d = parse_date("-".join(s.split('.')[::-1]))
            return d.isoformat() if d else None
        return _norm_month_date(s)

    picked_from_table = False
    for tbl in doc.tables:
        # собираем текст ячеек
        rows = []
        for r in tbl.rows:
            row_texts = []
            for c in r.cells:
                # в ячейке могут быть параграфы — склеим
                cell_txt = _clean("\n".join(p.text for p in c.paragraphs))
                row_texts.append(cell_txt)
            rows.append(row_texts)

        if not rows:
            continue

        # ищем заголовок таблицы
        header = [x.lower() for x in rows[0]]
        if len(header) >= 3 and \
                ('какого времени' in header[0]) and ('какое время' in header[1]):
            # ожидаем минимум 3 колонки
            for r in rows[1:]:
                if len(r) < 3:
                    continue
                f, t, pos = r[0].strip(), r[1].strip(), r[2].strip()
                if not f:
                    continue
                if not re.match(r'^\d{2}\.\d{4}$', f) and not re.match(r'^\d{2}\.\d{2}\.\d{4}$', f):
                    # это не строка данных
                    continue
                payload["service_history"].append({
                    "from": _norm_cell_date(f),
                    "to": _norm_cell_date(t),
                    "position": _clean(pos),
                })
            picked_from_table = True
            break  # нашли нужную таблицу — выходим

    # 6) Если таблицу не нашли, оставляем твой прежний «параграфный» бэкап
    if not picked_from_table:
        try:
            idx = para_lines.index("Самостоятельная трудовая деятельность и военная служба в ВС:")
        except ValueError:
            idx = -1
        if idx != -1:
            start = None
            for j in range(idx + 1, len(para_lines)):
                if re.match(r'^\d{2}\.\d{4}$', para_lines[j]) or re.match(r'^\d{2}\.\d{2}\.\d{4}$', para_lines[j]):
                    start = j
                    break
            if start is not None:
                k = start
                while k + 2 < len(para_lines):
                    f, t, pos = para_lines[k].strip(), para_lines[k + 1].strip(), para_lines[k + 2].strip()
                    if not (re.match(r'^\d{2}\.\d{4}$', f) or re.match(r'^\d{2}\.\d{2}\.\d{4}$', f)):
                        break
                    payload["service_history"].append({
                        "from": _norm_cell_date(f),
                        "to": _norm_cell_date(t),
                        "position": _clean(pos)
                    })
                    k += 3

    # блок подписи
    # три последние смысловые строки: должность HR, организация, "звание   Фамилия И."
    # найдём снизу
    tail = [ln for ln in para_lines[-10:] if ln]  # хвост документа
    if len(tail) >= 3:
        payload["sign_block"] = {
            "hr_title": tail[-3],
            "organization": tail[-2],
            "hr_rank": tail[-1].split()[0] if tail[-1] else None,
            "hr_name": " ".join(tail[-1].split()[1:]) if tail[-1] else None
        }

    return payload


class LD8ZipImportView(APIView):
    permission_classes = [IsAuthenticated, (IsHR | IsAdminOrRoot)]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        f = request.FILES.get("zip")
        if not f:
            return Response({"detail": "zip file is required"}, status=400)

        dry_run = str(request.data.get("dry_run", "true")).lower() in ("1", "true", "yes", "on")
        create_users = str(request.data.get("create_users", "false")).lower() in ("1", "true", "yes", "on")
        set_rank = str(request.data.get("set_rank", "false")).lower() in ("1", "true", "yes", "on")
        unit_id = request.data.get("unit_id")
        unit = None
        if unit_id:
            from apps.directory.models import Unit
            unit = get_object_or_404(Unit, pk=unit_id)

        try:
            zf = ZipFile(f)
        except BadZipFile:
            return Response({"detail": "Неверный ZIP архив"}, status=400)

        results = []
        errors = []

        for name in zf.namelist():
            if not name.lower().endswith(".docx"):
                continue
            try:
                content = zf.read(name)
                parsed = parse_ld8_docx(content, name)
                results.append(parsed)
            except Exception as e:
                errors.append({"file": name, "error": str(e)})

        timestamp = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
        out_dir = os.path.join(settings.MEDIA_ROOT, "imports")
        os.makedirs(out_dir, exist_ok=True)
        out_path = os.path.join(out_dir, f"ld8-{timestamp}.json")
        with open(out_path, "w", encoding="utf-8") as fp:
            json.dump({"parsed": results, "errors": errors}, fp, ensure_ascii=False, indent=2)
        out_url = settings.MEDIA_URL.rstrip("/") + "/imports/" + os.path.basename(out_path)

        if not create_users:
            return Response(
                {"saved_json": out_url, "created": [], "skipped": len(results), "errors": errors},
                status=200
            )

        # Создание записей (очень аккуратно, best-effort)
        created = []
        with transaction.atomic():
            from django.contrib.auth import get_user_model
            User = get_user_model()
            for item in results:
                email = (item.get("email") or "").strip().lower()
                if not email:
                    errors.append({"file": item.get("source_file"), "error": "email not found; skipped"})
                    continue
                # можно придумать правила генерации email, но пока оставим None — админ поправит позже
                user, created_user = User.objects.get_or_create(
                    email=email,
                    defaults={"role": getattr(User.UserRole, "OFFICER", "OFFICER")}
                )

                if created_user:
                    user.set_password("Testpass123")
                    user.save(update_fields=["password"])

                prof, _ = OfficerProfile.objects.get_or_create(user=user)
                fio = item.get("full_name") or ""
                if fio:
                    prof.full_name = fio

                b = item.get("birth") or {}
                if b.get("date"):
                    prof.birth_date = parse_date(b["date"])
                if b.get("place"):
                    prof.birth_place = b["place"]

                if item.get("iin"):
                    prof.iin = item["iin"]
                if item.get("nationality"):
                    prof.nationality = item["nationality"]

                # Семейное положение (нормализация строки в choices)
                ms_raw = item.get("marital_status") or ""
                ms_norm = _map_marital_status(ms_raw)
                if ms_norm:
                    prof.marital_status = ms_norm  # SINGLE/MARRIED/DIVORCED/WIDOWED

                if unit:
                    prof.unit = unit

                # Участие в боевых действиях
                cp_text = (item.get("combat_participation") or "").strip()
                cp_flag = cp_text != "" and any(x in cp_text.lower() for x in ("участник", "боев", "миротвор"))
                prof.combat_participation = cp_flag
                if cp_text:
                    prof.combat_notes = cp_text  # <-- ДОБАВИЛИ: сохраняем исходную строку как заметку

                # звание
                if set_rank and item.get("rank", {}).get("name"):
                    try:
                        prof.rank = Rank.objects.get(name__iexact=item["rank"]["name"].strip())
                    except Rank.DoesNotExist:
                        pass
                    # дату начала службы можно прибить к service_start_date, если пусто
                    if not prof.service_start_date and item["rank"].get("since"):
                        prof.service_start_date = parse_date(item["rank"]["since"])
                prof.save()
                created.append({"user_id": user.id, "profile_id": prof.id, "full_name": prof.full_name})

        return Response({"dry_run": False, "created": created, "parsed_count": len(results), "errors": errors},
                        status=201)
